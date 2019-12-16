def Lebedev():
    from words_processing import words_processing

    import glossary
    txt = glossary.Text()

    import re
    from docx import Document
    FileName = "D:\\Programs\\myGlossary\\data\\dict\\Лебедев.docx"
    f = open(FileName, 'rb')
    txt.text = f.read()

    document = Document(f)
    sentencesInfo = []
    for i, paragraph in enumerate(document.paragraphs):
        print("%d предложение из %d" % (i + 1, len(document.paragraphs)))

        S = paragraph.text.find("–")
        if S != -1:
            line = paragraph.text.split("–")[::-1]
            eng0 = re.findall("[A-Za-z]", line[0])
            rus0 = re.findall("[А-Яа-я]", line[0])
            eng1 = re.findall("[A-Za-z]", line[1])
            rus1 = re.findall("[А-Яа-я]", line[1])
            if (len(rus0) == 0 ) and (len(eng1) == 0):
                print(line)
                words0 = words_processing.sentencesToWords(line[0])
                wordsInfo0 = []
                for word in words0:
                    normalWord = words_processing.getEnglishNormalForm(word)
                    wordsInfo0.append({"currentWord": word, "normalWord": normalWord, "isMain": True})

                words1 = words_processing.sentencesToWords(line[1])
                wordsInfo1 = []
                for word in words1:
                    normalWord = words_processing.getRussianNormalForm(word)
                    wordsInfo1.append({"currentWord": word, "normalWord": normalWord, "isMain": True})

                if (len(wordsInfo0) > 0) and (len(wordsInfo1) > 0):
                    sentencesInfo.append({"pairs":
                                              [{"sentence": line[0], "wordsInfo": wordsInfo0, "lineOfWords": None},
                                               {"sentence": line[1], "wordsInfo": wordsInfo1, "lineOfWords": None}],
                                          "correspondences": [],
                                          "tags": ["научный словарь Лебедева"]})

    txt.textInfo = [{"paragraph": txt.text, "sentencesInfo": sentencesInfo}]
    import pickle
    f = open("D:\\Programs\\myGlossary\\data\\dict\\Лебедев.gls", "wb")
    pickle.dump(txt, f)


def Tsimmerman():
    from words_processing import words_processing

    import glossary
    txt = glossary.Text()

    import re
    from docx import Document
    FileName = "D:\\Programs\\myGlossary\\data\\dict\\Циммерман.docx"
    f = open(FileName, 'rb')
    txt.text = f.read()

    document = Document(f)
    sentencesInfo = []
    currentRus = None
    currentEng = None
    for i, paragraph in enumerate(document.paragraphs):
        # if (i == 100): break

        print("%d предложение из %d" % (i + 1, len(document.paragraphs)))
        line = paragraph.text.strip()
        if line == "": continue

        eng = re.findall("[A-Za-z]", line)
        rus = re.findall("[А-Яа-я]", line)
        IsEnglish = len(eng) > len(rus)

        if not IsEnglish:
            currentRus = line
        else:
            currentEng = line

            if (currentRus is not None) and (currentEng is not None):
                print(currentRus, currentEng)
                words0 = words_processing.sentencesToWords(currentEng)
                wordsInfo0 = []
                for word in words0:
                    normalWord = words_processing.getEnglishNormalForm(word)
                    wordsInfo0.append({"currentWord": word, "normalWord": normalWord, "isMain": True})

                words1 = words_processing.sentencesToWords(currentRus)
                wordsInfo1 = []
                for word in words1:
                    normalWord = words_processing.getRussianNormalForm(word)
                    wordsInfo1.append({"currentWord": word, "normalWord": normalWord, "isMain": True})

                if (len(wordsInfo0) > 0) and (len(wordsInfo1) > 0):
                    sentencesInfo.append({"pairs":
                                              [{"sentence": currentEng, "wordsInfo": wordsInfo0, "lineOfWords": None},
                                               {"sentence": currentRus, "wordsInfo": wordsInfo1, "lineOfWords": None}],
                                          "correspondences": [],
                                          "tags": ["научный словарь Циммермана"]})

    txt.textInfo = [{"paragraph": txt.text, "sentencesInfo": sentencesInfo}]
    # txt.printData()

    import pickle
    f = open("D:\\Programs\\myGlossary\\data\\dict\\Циммерман.gls", "wb")
    pickle.dump(txt, f)


def Andreeva():
    import glossary
    txt = glossary.Text()
    txt.authors = "Андреева"
    txt.caption = "Статьи"
    txt.year = "2000"
    txt.data = ""

    import re
    FileNamePattern = "D:\\Programs\\myGlossary\\data_noGPR\\Andreeva_papers"
    InputFileName = FileNamePattern + ".txt"
    OutputFileName = FileNamePattern + ".gls"

    f = open(InputFileName, 'r', encoding="utf-8")
    sourceText = f.read()
    lines = sourceText.split("\n")

    data_english = []
    data_russian = []

    eng = len(re.findall("[A-Za-z]", lines[0]))
    rus = len(re.findall("[А-Яа-я]", lines[0]))

    if rus > eng:
        f = False
    else:
        f = True

    ok = True
    for line in lines:
        eng = len(re.findall("[A-Za-z]", line))
        rus = len(re.findall("[А-Яа-я]", line))
        if f:
            if eng > rus:
                data_english.append(line)
                f = not f
            else:
                print("Подряд идет русский текст! \n" + line)
                ok = False
        else:
            if rus > eng:
                data_russian.append(line)
                f = not f
            else:
                print("Подряд идет английский текст! \n" + line)
                ok = False

    if not ok: return



    if len(data_english) != len(data_russian):
        print("Количество русских предложений не совпадает с количеством английских!")
        return

    txt.text = "\n".join(data_english)

    txt.initOneParagraph(len(data_english))


    for i in range(len(data_english)):
        txt.setSentence(0, i, data_english[i])
        txt.setTranslation(0, i, data_russian[i])
    # txt.printData()


    txt.calculate()

    import pickle
    f = open(OutputFileName, "wb")
    pickle.dump(txt, f)




    # document = Document(f)
    # sentencesInfo = []
    # for i, paragraph in enumerate(document.paragraphs):
    #     print("%d предложение из %d" % (i + 1, len(document.paragraphs)))
    #
    #     S = paragraph.text.find("–")
    #     if S != -1:
    #         line = paragraph.text.split("–")[::-1]
    #         eng0 = re.findall("[A-Za-z]", line[0])
    #         rus0 = re.findall("[А-Яа-я]", line[0])
    #         eng1 = re.findall("[A-Za-z]", line[1])
    #         rus1 = re.findall("[А-Яа-я]", line[1])
    #         if (len(rus0) == 0) and (len(eng1) == 0):
    #             print(line)
    #             words0 = words_processing.sentencesToWords(line[0])
    #             wordsInfo0 = []
    #             for word in words0:
    #                 normalWord = words_processing.getEnglishNormalForm(word)
    #                 wordsInfo0.append({"currentWord": word, "normalWord": normalWord, "isMain": True})
    #
    #             words1 = words_processing.sentencesToWords(line[1])
    #             wordsInfo1 = []
    #             for word in words1:
    #                 normalWord = words_processing.getRussianNormalForm(word)
    #                 wordsInfo1.append({"currentWord": word, "normalWord": normalWord, "isMain": True})
    #
    #             if (len(wordsInfo0) > 0) and (len(wordsInfo1) > 0):
    #                 sentencesInfo.append({"pairs":
    #                                           [{"sentence": line[0], "wordsInfo": wordsInfo0, "lineOfWords": None},
    #                                            {"sentence": line[1], "wordsInfo": wordsInfo1, "lineOfWords": None}],
    #                                       "correspondences": [],
    #                                       "tags": ["научный словарь Лебедева"]})
    #
    # txt.textInfo = [{"paragraph": txt.text, "sentencesInfo": sentencesInfo}]
    # import pickle
    # f = open("D:\\Programs\\myGlossary\\data\\dict\\Лебедев.gls", "wb")
    # pickle.dump(txt, f)


if __name__ == '__main__':
    Andreeva()
    # Lebedev()
    # Tsimmerman()